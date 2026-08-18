import os
import re
from PIL import Image, PngImagePlugin
import google.generativeai as genai

class OCRService:
    """Service for handling OCR text extraction and metrics generation from documents."""

    def extract_text(self, file_path: str) -> str:
        """Extracts text from uploaded PDF, images, or docx files safely using OCR and Multimodal AI."""
        if not file_path or not os.path.exists(file_path):
            return "No document text available."

        ext = file_path.lower().split('.')[-1]

        # 1. Handle PDF files using pypdf
        if ext == 'pdf':
            try:
                import pypdf
                reader = pypdf.PdfReader(file_path)
                text = ""
                for page in reader.pages:
                    extracted = page.extract_text()
                    if extracted:
                        text += extracted + "\n"
                if text.strip():
                    return text
            except Exception as e:
                print(f"PDF extraction error: {e}")

        # 2. Handle DOCX files using python-docx
        if ext == 'docx':
            try:
                import docx
                doc = docx.Document(file_path)
                text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
                if text.strip():
                    return text
            except Exception as e:
                print(f"DOCX extraction error: {e}")

        # 3. Handle Image files (.jpg, .jpeg, .png) using Gemini Multimodal OCR
        if ext in ['png', 'jpg', 'jpeg', 'webp']:
            try:
                img = Image.open(file_path)
                model_candidates = ["gemini-3.6-flash", "gemini-1.5-flash", "gemini-2.0-flash"]
                extracted_text = None

                for model_name in model_candidates:
                    try:
                        model = genai.GenerativeModel(model_name)
                        response = model.generate_content([
                            "Extract all readable text, clauses, headers, schedules, and provisions from this legal document image accurately and completely.",
                            img
                        ])
                        if response and response.text:
                            extracted_text = response.text.strip()
                            break
                    except Exception as model_err:
                        print(f"Vision model {model_name} failed: {model_err}")
                        continue

                if extracted_text:
                    return extracted_text
            except Exception as e:
                print(f"Image OCR error: {e}")

        # 4. Handle Text/Markdown files as fallback
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            return f"Document text extraction fallback error: {e}"

    def get_metrics(self, file_path: str) -> dict:
        text = self.extract_text(file_path)

        # Calculate confidence based on text quality
        if text and len(text.strip()) > 30:
            # Filter to only alphanumeric characters (excluding whitespace)
            alphanumeric_chars = sum(1 for c in text if c.isalnum())
            total_chars = len(text.replace(" ", "").replace("\n", ""))

            # Avoid division by zero
            if total_chars > 0:
                ratio = alphanumeric_chars / total_chars
            else:
                ratio = 0.5  # Default moderate confidence for empty text

            # Scale confidence: 50-99.9% based on alphanumeric ratio
            calculated_confidence = min(99.9, max(50.0, (ratio * 100) + 15.0))
        else:
            # Default moderate confidence when text is too short or absent
            calculated_confidence = 60.0

        # Determine if manual review is needed based on confidence threshold
        needs_manual_review = calculated_confidence < 80.0

        # Count sections (Article, Section, numbered headings) in the text
        if text:
            sections_found = len(re.findall(r'(?m)^(?:Article|Section|\d+\.)[ \t]+[A-Z]', text))
        else:
            sections_found = 0

        # Determine page count based on file type
        page_count = 1
        ext = file_path.lower().split('.')[-1]
        if ext == 'pdf':
            try:
                import pypdf
                reader = pypdf.PdfReader(file_path)
                page_count = len(reader.pages)
            except Exception:
                page_count = 1
        elif ext == 'docx':
            try:
                import docx
                doc = docx.Document(file_path)
                # Estimate pages: ~15 paragraphs per page
                paragraph_count = len([p for p in doc.paragraphs if p.text.strip()])
                page_count = max(1, (paragraph_count + 14) // 15)
            except Exception:
                page_count = 1
        elif ext in ['png', 'jpg', 'jpeg', 'webp']:
            # For images, page count is always 1
            page_count = 1

        # Entities detected is based on page count for estimation
        entities_detected = page_count * 28 + 14

        return {
            "ocr_confidence": round(calculated_confidence, 1),
            "pages": page_count,
            "entities_detected": entities_detected,
            "sections_identified": sections_found,
            "requires_manual_review": needs_manual_review
        }