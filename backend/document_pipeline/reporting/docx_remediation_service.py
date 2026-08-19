import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

class DocxRemediationService:
    def generate_schedule_of_deviations(self, doc_data: dict, clauses: list, output_path: str) -> str:
        doc = Document()
        
        # Title Section
        title = doc.add_heading(level=1)
        run = title.add_run("Tata Group - Schedule of Contract Deviations & AI Redlines")
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0, 43, 73)  # Tata Deep Navy
        
        subtitle = doc.add_paragraph(f"Job ID: {doc_data.get('job_id', 'N/A')} | Filename: {doc_data.get('filename', 'Contract.docx')}")
        subtitle.runs[0].font.size = Pt(10)
        subtitle.runs[0].font.color.rgb = RGBColor(100, 116, 139)
        
        doc.add_paragraph()

        # Add Deviations Table
        table = doc.add_table(rows=1, cols=4)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        headers = ["Clause Type", "Original Text", "Risk Rationale & Policy", "AI Proposed Redline"]
        for i, header_text in enumerate(headers):
            hdr_cells[i].text = header_text
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(255, 255, 255)
            # Set header background color
            shading_elm = parse_xml(r'<w:shd {} w:fill="002B49"/>'.format(nsdecls('w')))
            hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)

        # Populate rows with extracted clauses and proposed redlines
        for clause in clauses:
            if not clause:
                continue
            row_cells = table.add_row().cells
            row_cells[0].text = str(clause.get('clause_type', 'General'))
            row_cells[1].text = str(clause.get('extracted_text', ''))
            row_cells[2].text = f"Risk: {clause.get('risk_level', 'LOW')}\nRationale: {clause.get('risk_rationale', '')}\nRef: {clause.get('rag_reference_used', 'N/A')}"
            row_cells[3].text = str(clause.get('proposed_redline') or 'No redline required (Standard Compliance)')

            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(8.5)

        doc.save(output_path)
        return output_path
